"""
docs_tool.py — Generate documents and save to the Generated Docs folder.
Supports PDF, Word, text, markdown, HTML, CSV.
Auto-opens the file after saving.
"""
import os
import re
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import GENERATED_DOCS_DIR

DEFINITIONS = [
    {
        "name": "generate_document",
        "description": (
            "Generate a document and save it to the Generated Docs folder. "
            "Use this for any output that isn't going directly into a codebase: "
            "reports, invoices, summaries, plans, research docs, data exports, proposals — "
            "anything Ky might want to open and read. "
            "The file is saved and opened automatically."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Filename without directory path, e.g. 'invoice-june-2026.pdf' or 'market-research.docx'",
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Document content in markdown. Use # for headings, **bold**, - for lists, "
                        "| for tables. For CSV format use comma-separated rows."
                    ),
                },
                "format": {
                    "type": "string",
                    "enum": ["pdf", "docx", "txt", "md", "html", "csv"],
                    "description": "Output format. Use pdf or docx for polished documents, md for notes, csv for data.",
                },
            },
            "required": ["filename", "content", "format"],
        },
    }
]


def handle(name: str, inputs: dict) -> str:
    if name == "generate_document":
        return _generate(
            inputs.get("filename", "document.txt"),
            inputs.get("content", ""),
            inputs.get("format", "txt"),
        )
    return f"Unknown tool: {name}"


def _ensure_dir() -> str:
    os.makedirs(GENERATED_DOCS_DIR, exist_ok=True)
    return GENERATED_DOCS_DIR


def _generate(filename: str, content: str, fmt: str) -> str:
    folder = _ensure_dir()
    base   = os.path.splitext(filename)[0]
    path   = os.path.join(folder, f"{base}.{fmt}")

    try:
        if fmt in ("txt", "md", "html", "csv"):
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
        elif fmt == "docx":
            _write_docx(path, content)
        elif fmt == "pdf":
            _write_pdf(path, content)
        else:
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
    except ImportError as e:
        return f"Missing library for {fmt} generation: {e}. Install it with pip."
    except Exception as e:
        return f"Failed to generate {fmt}: {e}"

    # Auto-open with default app
    try:
        os.startfile(path)
    except Exception:
        pass

    return f"Saved to {path}"


def _write_docx(path: str, content: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        elif stripped == "":
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            # Handle inline **bold** and *italic*
            parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)
    doc.save(path)


def _write_pdf(path: str, content: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    doc  = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=25*mm, rightMargin=25*mm,
        topMargin=25*mm, bottomMargin=25*mm,
    )
    styles = getSampleStyleSheet()

    # Clean up bullet/list style to avoid missing style errors
    body   = ParagraphStyle("body",   parent=styles["Normal"],  fontSize=11, leading=16)
    bullet = ParagraphStyle("bullet", parent=styles["Normal"],  fontSize=11, leading=16, leftIndent=12)
    h1     = ParagraphStyle("h1",     parent=styles["Heading1"], fontSize=18, spaceAfter=6)
    h2     = ParagraphStyle("h2",     parent=styles["Heading2"], fontSize=14, spaceAfter=4)
    h3     = ParagraphStyle("h3",     parent=styles["Heading3"], fontSize=12, spaceAfter=3)

    story = []
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):
            story.append(Paragraph(stripped[4:], h3))
        elif stripped.startswith("## "):
            story.append(Paragraph(stripped[3:], h2))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], h1))
            story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc"), spaceAfter=4))
        elif stripped.startswith(("- ", "* ")):
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', stripped[2:])
            story.append(Paragraph(f"• {text}", bullet))
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', re.sub(r"^\d+\. ", "", stripped))
            story.append(Paragraph(text, bullet))
        elif stripped == "":
            story.append(Spacer(1, 8))
        else:
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', stripped)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            story.append(Paragraph(text, body))

    doc.build(story)
